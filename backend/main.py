# importing libraries
from fastapi import FastAPI, HTTPException, Response, status, UploadFile, Request, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi import BackgroundTasks
from PyPDF2 import PdfReader
from urllib.parse import urlparse
from PIL import Image
import json
import base64
import asyncio
import textract
import os
import pytesseract
import cv2
import numpy as np
import time
import docx
# importing in-house code
from model.api import API
import model.data_structures as ds
# import model.social_media_scraper as sms
# import model.generic_web_scraper as gws
import model.web_scrapers as ws
import model.graph as graph
import model.gauge as gauge
import model.text as text_utils
import model.validation as validation

# Initialising constants 
# setting file paths for configuration files
CONFIG_FILE_PATH = r"config\config.json"
TEMP_FILE_PATH = r"config\temp.json"
API_FILE_PATH = r"config\api.json"
RESULTS_FILE_PATH = r"config\results.json"
SCORE_FILE_PATH = r"config\scores.json"
# setting file path to tesseract
pytesseract.pytesseract.tesseract_cmd = r"Tesseract-OCR\tesseract.exe"

# set fastapi up
botbuster = FastAPI()

# set up fastapi to bypass CORS at the front end:
botbuster.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# writing endpoints
@botbuster.post("/splittext/")
def split_text(request: ds.split_text):
    list_of_apis = request.dict()["list_of_apis"]
    full_text = request.dict()["text"]
    # validating full_text
    if(validation.check_text(full_text)==False):
        raise HTTPException(status_code=400, detail="Bad Request")
    with open(CONFIG_FILE_PATH, "r") as config_data_file: # load in config data from config file
        config_data = json.load(config_data_file)
    list_of_texts = text_utils.chunk(full_text, config_data["chunk_option"])
    with open(SCORE_FILE_PATH, "w") as score_data_file: # empty in score file
        scores = {"overall_score": {api_category: {} for api, api_category in list_of_apis}}
        scores["overall_score"]["sentence_data"] = {}
        scores["overall_score"]["sentence_data"]["total_num_sentences"] = 0
        score_data_file.write(json.dumps(scores, indent=4))
    return list_of_texts


# calling apis to check the text
@botbuster.post("/checktext/new") # endpoint #1 sending requests to the AI detection engines
def check_text(request: ds.check_text):
    decoded_req=validation.decode(request.dict())# decode the encoded text
    list_of_apis = decoded_req["list_of_apis"]
    text = decoded_req["text"]
    # validate text
    if(validation.check_text(text)==False):
        raise HTTPException(status_code=400, detail="Bad Request")
    list_of_apis = request.dict()["list_of_apis"]
    text = request.dict()["text"]
    with open(SCORE_FILE_PATH, "r") as score_data_file: # load in config data from config file
        scores = json.load(score_data_file)
    full_results = {api : {} for api, api_category in list_of_apis} # create dictionary to store all results
    total_score = {api_category:{} for api, api_category in list_of_apis} # store score of each API before taking the average
    num_sentences = text_utils.chunk_size(text)
    request_num = len(scores)
    # print(f"Request No.: {request_num}")
    scores["overall_score"]["sentence_data"][request_num] = num_sentences
    scores["overall_score"]["sentence_data"]["total_num_sentences"] += num_sentences
    with open(CONFIG_FILE_PATH, "r") as config_data_file: # load in config data from config file
        config_data = json.load(config_data_file)
    seen_categories = [] 
    for api_num, [api, api_category] in enumerate(list_of_apis): # loop through all APIs
        scores["overall_score"][api_category][api] = "score not calculated"
        try:
            if api != "score_type" and api != "description":
                results = API(config_data["APIs"][api_category][api]).api_call(text)                
                full_results[api] = results
            if api_category not in seen_categories:
                total_score[api_category] = {"score": 0, "num_apis": 0}
            if api_num == 0:
                scores[request_num] = {
                    "general_score": {api_category : {
                        "description": config_data["APIs"][api_category]["description"],
                        "score_type": config_data["APIs"][api_category]["score_type"]
                    } for api, api_category in list_of_apis},
                    "sentence_score": []  # keeping a score for each sentence  
                }
                for sentence in text_utils.chunk_by_sentences(text):
                    scores[request_num]["sentence_score"].append({sentence: {}})
            seen_categories.append(api_category)
        except Exception:
            full_results[api] = "Error Detecting"
            continue
        try:
            # checking for general score 
            results = full_results # set the results to loop through
            if results == "Error Detecting":
                scores[request_num]["general_score"][api_category][api] = "error detecting" 
            path = config_data["path_to_general_score"][api] # get the path
            for key in path.split('.'): # loop through each key in the path to general score
                if str(key).isnumeric(): # if the key is numerical, convert it to an int
                    key = int(key)
                if key == "num":
                    key = request_num
                try: 
                    results = results[key] # try to path to the score
                    if config_data["APIs"][api_category]["score_type"] == "Discrete" and key == path.split('.')[-1]:
                        if results == "flag":
                            results = 1
                        else:
                            results = 0
                    score = round(float(results) * 100,1)
                    if score > 100:
                        scores[request_num]["general_score"][api_category][api] = "error getting score"
                        continue
                    scores[request_num]["general_score"][api_category][api] = score # appends general score to the dictionary
                    total_score[api_category]["score"] += score # sums the general scores of all apis
                    total_score[api_category]["num_apis"] += 1 # sums the number of APIs in the category
                    final_score = total_score[api_category]["score"]/total_score[api_category]["num_apis"] # gets the average score of all the APIs in each category
                    scores[request_num]["general_score"][api_category]["overall_score"] = round(final_score,1) # appends the average score
                except TypeError: # Type error will occur if the loop hasn't reached the score
                    continue  # continue to the next key in the loop
                except KeyError: # If there is an error with the path
                    scores[request_num]["general_score"][api_category][api] = "error getting score" 
                    continue
                except Exception: # catch all other errors
                    continue      
        except Exception:
            continue  
    try:
        scores[request_num]["flags"] = []
        for api_category in scores[request_num]["general_score"]:
            try:
                if scores[request_num]["general_score"][api_category]["overall_score"] > config_data["flagged_threshold"]:
                    scores[request_num]["flags"].append(f"Flagged by ${api_category}")
                    continue
                if scores[request_num]["general_score"][api_category]["overall_score"] > config_data["potentially_flagged_threshold"]:
                    scores[request_num]["flags"].append(f"Potentially Flagged by {api_category}")
            except Exception:
                continue   
    except Exception:
        pass
    # checking for sentence score
    for api, api_category in list_of_apis:
        try:
            path1 = config_data["path_to_sentence_score"][api][0] # path to the list of the sentences
            path2 = config_data["path_to_sentence_score"][api][1] # path to the score of each sentence
            results = full_results # checking for sentence score
            for key in path1.split("."): # loops through each key in the path to sentence scores
                try:
                    if key.isnumeric(): # if the key is numerical, convert it to an int
                        key = int(key)
                    if key == "num":
                        key = request_num
                    results = results[key] # try to path to the score
                except KeyError or TypeError or Exception:
                    continue
            for num in range(0, len(results)): # loop through each sentence of results
                try: 
                    for sentence in scores[request_num]["sentence_score"][num].keys():
                        if not(api_category in scores[request_num]["sentence_score"][num][sentence]):
                            scores[request_num]["sentence_score"][num][sentence] = {
                                api_category: [],
                                f"{api_category}-highlight": 0
                            }
                        if sentence == results[num]["sentence"] and results[num][path2] > 0.70: # if score above > 70%, add the  highlight and the api name 
                            scores[request_num]["sentence_score"][num][sentence][api_category].append(api)
                            scores[request_num]["sentence_score"][num][sentence][f"{api_category}-highlight"] += 1/total_score[api_category]["num_apis"]
                except:
                    continue  
        except KeyError: # Key Error will trigger is there is no path to sentence score (API does not have this capability)
            continue      
    for api_category in scores["overall_score"].keys():
        try:
            if api_category == "sentence_data":
                continue
            api_count = 0
            api_category_total_score = 0
            if request_num == 1:
                keys = scores[1]["general_score"][api_category]
            else: 
                keys = scores["1"]["general_score"][api_category]
            for api in keys: 
                if api == "description" or api == "score_type" or api == "overall_score":
                    continue
                if api not in scores["overall_score"][api_category]:
                    scores["overall_score"][api_category][api] = "score not calculated"
                    continue
                api_total_score = 0
                for req_num in scores.keys():
                    if req_num == "overall_score":
                        continue
                    try: 
                        api_total_score += scores[req_num]["general_score"][api_category][api]  * (scores["overall_score"]["sentence_data"][req_num]/scores["overall_score"]["sentence_data"]["total_num_sentences"])
                    except Exception:
                        continue
                if str(int(api_total_score)).isnumeric() and scores[req_num]["general_score"][api_category][api] != "error getting score":
                    scores["overall_score"][api_category][api] = round(api_total_score,1)
                    api_count += 1
                    api_category_total_score += api_total_score
                scores["overall_score"][api_category]["average_score"] = round(api_category_total_score/api_count, 1)
        except:
            scores["overall_score"][api_category][api] = "score not calculated"
    # with open(RESULTS_FILE_PATH, "w") as results_file: # load in API data from api file
    #     results_file.write(json.dumps(full_results, indent=4))
    with open(SCORE_FILE_PATH, "w") as score_data_file: # load in API data from api file
        score_data_file.write(json.dumps(scores, indent=4))
    graph.generate_graph(scores["overall_score"]) # generate the graph with the overall scores of each API
    gauge.generate_gauge(scores["overall_score"]) # generate the gauge with the overall scores of each API
    return validation.sanitise(scores)

# calling apis to check the text
@botbuster.post("/checktext/") # endpoint #1 sending requests to the AI detection engines
def check_text(request: ds.check_text):
    decoded_req=validation.decode(request.dict())
    list_of_apis = decoded_req["list_of_apis"]
    full_text = decoded_req["text"]
    full_results = {api : {} for api, api_category in list_of_apis} # create dictionary to store all results
    scores = {} # create dictionary to store all scores
    overall_scores = {api_category: {} for api, api_category in list_of_apis}
    total_score = {api_category:{} for api, api_category in list_of_apis} # store score of each API before taking the average
    overall_scores["sentence_data"] = {}
    total_num_sentences = text_utils.chunk_size(full_text)
    overall_scores["sentence_data"]["total_num_sentences"] = total_num_sentences
    with open(CONFIG_FILE_PATH, "r") as config_data_file: # load in config data from config file
        config_data = json.load(config_data_file)
    # change this to the text chunking functions
    list_of_texts = text_utils.chunk(full_text, config_data["chunk_option"])  
    seen_categories = [] 
    for api_num, [api, api_category] in enumerate(list_of_apis): # loop through all APIs
        overall_scores[api_category][api] = "score not calculated"
        try:
            # full_results = {}
            for num, text in enumerate(list_of_texts):
                if api != "score_type" and api != "description":
                    results = API(config_data["APIs"][api_category][api]).api_call(text[0])
                    full_results[api][num] = results
                    num_sentences = text_utils.chunk_size(text[0])
                    overall_scores["sentence_data"][num + 1] = num_sentences
                if api_category not in seen_categories:
                    total_score[api_category][num + 1] = {"score": 0, "num_apis": 0}
                if api_num == 0:
                    scores[num + 1] = {
                        "general_score": {api_category : {
                            "description": config_data["APIs"][api_category]["description"],
                            "score_type": config_data["APIs"][api_category]["score_type"]
                        } for api, api_category in list_of_apis},
                        "sentence_score": []  # keeping a score for each sentence  
                    }
                    for sentence in text_utils.chunk_by_sentences(text[0]):
                        scores[num + 1]["sentence_score"].append({sentence: {}})
            
            seen_categories.append(api_category)
        except Exception:
            full_results[api][num] = "Error Detecting"
            continue
        try:
            # checking for general score 
            for req_num in full_results[api].keys():
                results = full_results # set the results to loop through
                if results == "Error Detecting":
                    scores[req_num + 1]["general_score"][api_category][api] = "error detecting" 
                path = config_data["path_to_general_score"][api] # get the path
                for key in path.split('.'): # loop through each key in the path to general score
                    if str(key).isnumeric(): # if the key is numerical, convert it to an int
                        key = int(key)
                    if key == "num":
                        key = req_num
                    try: 
                        results = results[key] # try to path to the score
                        if config_data["APIs"][api_category]["score_type"] == "Discrete" and key == path.split('.')[-1]:
                            if results == "flag":
                                results = 1
                            else:
                                results = 0
                        scores[req_num + 1]["general_score"][api_category][api] = round(float(results) * 100,1) # appends general score to the dictionary
                        total_score[api_category][req_num + 1]["score"] += round(float(results) * 100,1) # sums the general scores of all apis
                        total_score[api_category][req_num + 1]["num_apis"] += 1 # sums the number of APIs in the category
                        final_score = total_score[api_category][req_num + 1]["score"]/total_score[api_category][req_num + 1]["num_apis"] # gets the average score of all the APIs in each category
                        scores[req_num + 1]["general_score"][api_category]["overall_score"] = round(final_score,1) # appends the average score
                        overall_scores[api_category][api][req_num + 1]["score"] = round(float(results) * 100,1)
                        
                    except TypeError: # Type error will occur if the loop hasn't reached the score
                        continue  # continue to the next key in the loop
                    except KeyError: # If there is an error with the path
                        scores[req_num + 1]["general_score"][api_category][api] = "error getting score" 
                        continue
                    except Exception: # catch all other errors
                        continue      
        except Exception:
            continue
    for req_num in scores.keys():
        if not str(req_num).isnumeric():
            continue
        scores[req_num]["flags"] = []
        try:
            for api_category in scores[req_num]["general_score"]:
                try:
                    if scores[req_num]["general_score"][api_category]["overall_score"] > config_data["flagged_threshold"]:
                        scores[req_num]["flags"].append(f"Flagged by ${api_category}")
                        continue
                    if scores[req_num]["general_score"][api_category]["overall_score"] > config_data["potentially_flagged_threshold"]:
                        scores[req_num]["flags"].append(f"Potentially Flagged by {api_category}")
                except Exception:
                    continue   
        except Exception:
            continue
        # checking for sentence score
        for api, api_category in list_of_apis:
            try:
                path1 = config_data["path_to_sentence_score"][api][0] # path to the list of the sentences
                path2 = config_data["path_to_sentence_score"][api][1] # path to the score of each sentence
                results = full_results # checking for sentence score
                for key in path1.split("."): # loops through each key in the path to sentence scores
                    try:
                        if key.isnumeric(): # if the key is numerical, convert it to an int
                            key = int(key)
                        if key == "num":
                            key = req_num - 1
                        results = results[key] # try to path to the score
                    except KeyError or TypeError or Exception:
                        continue
                for num in range(0, len(results)): # loop through each sentence of results
                    try: 
                        for sentence in scores[req_num]["sentence_score"][num].keys():
                            if not(api_category in scores[req_num]["sentence_score"][num][sentence]):
                                scores[req_num]["sentence_score"][num][sentence] = {
                                    api_category: [],
                                    f"{api_category}-highlight": 0
                                }
                            if sentence == results[num]["sentence"] and results[num][path2] > 0.70: # if score above > 70%, add the  highlight and the api name
                                scores[req_num]["sentence_score"][num][sentence][api_category].append(api)
                                scores[req_num]["sentence_score"][num][sentence][f"{api_category}-highlight"] += 1/total_score[api_category][req_num]["num_apis"]
                    except:
                        continue  
            except KeyError: # Key Error will trigger is there is no path to sentence score (API does not have this capability)
                continue       
    try:
        for api_category in overall_scores.keys():
            if api_category == "sentence_data":
                continue
            api_count = 0
            api_category_total_score = 0
            for api in scores[1]["general_score"][api_category]: 
                if api == "description" or api == "score_type" or api == "overall_score":
                    continue
                if api not in overall_scores[api_category]:
                    overall_scores[api_category][api] = "score not calculated"
                    continue
                api_total_score = 0
                for req_num in scores.keys():
                    if req_num == "overall_score":
                        continue
                    try: 
                        api_total_score += scores[req_num]["general_score"][api_category][api] * (overall_scores["sentence_data"][req_num]/overall_scores["sentence_data"]["total_num_sentences"])
                    except Exception:
                        continue
                if str(int(api_total_score)).isnumeric() and scores[req_num]["general_score"][api_category][api] != "error getting score":
                    overall_scores[api_category][api] = round(api_total_score,1)
                    api_count += 1
                    api_category_total_score += api_total_score
                overall_scores[api_category]["average_score"] = round(api_category_total_score/api_count, 1)
    except:
        overall_scores[api_category][api] = "score not calculated"
    # with open(RESULTS_FILE_PATH, "w") as results_file: # load in API data from api file
    #     results_file.write(json.dumps(full_results, indent=4))
    with open(r"config\score.json", "w") as scores_file: # load in API data from api file
        scores["overall_score"] = overall_scores
        scores_file.write(json.dumps(scores, indent=4))
    graph.generate_graph(scores["overall_score"]) # generate the graph with the overall scores of each API
    gauge.generate_gauge(scores["overall_score"]) # generate the gauge with the overall scores of each API
    scores["overall_score"] = overall_scores
    return validation.sanitise(scores)

@botbuster.get("/getapis/") # endpoint #2 retrieving available API information
async def get_apis():
    api_info = {} 
    try: 
        with open(API_FILE_PATH, "r") as api_data_file:
            api_data = json.load(api_data_file) # grabs all API data from API config file
            for api_category in api_data.keys(): # loops through the API categories
                api_info[api_category] = {} # creates an empty dictionary for each category
                
                for api in api_data[api_category].keys(): # loops through all API in each category 
                    api_info[api_category][api] = [api.lower().replace(" ", "")] # creates 
                    try: 
                        with open(CONFIG_FILE_PATH, "r") as config_data_file:
                            config_data = json.load(config_data_file)
                            api_logo_path = config_data["logos_base_path"] + config_data["logos"][api] # sets file path for the system to open
                            with open(api_logo_path, "rb") as image:
                                api_info[api_category][api].append(base64.b64encode(image.read()).decode("utf-8)")) # encodes the image in base 64 and decodes in utf-8
                    except Exception: 
                        continue # catches all errors and continues to the next api in the loop. If there is no logo etc, it will just not show up on UI
    except Exception:
        raise HTTPException (status_code = 500, detail = "Internal Server Error")
    return api_info

# adding apis to system
@botbuster.post("/addapi/") # endpoint #3 adding APIs to the system
def add_api(request: ds.add_api, response: Response):
    req = request.dict()
    try:
        API(req["api_details"])
        api_name = req["api_name"]
    except KeyError:
        raise HTTPException (status_code = 400, detail = "missing details")
    else:
        with open(API_FILE_PATH, "r") as add_api_file:
            api_data = json.load(add_api_file)
            api_data[api_name] = req["api_details"]
        with open(API_FILE_PATH, "w") as add_api_file:
            add_api_file.write(json.dumps(api_data, indent = 4))
            response.status_code = status.HTTP_204_NO_CONTENT

# scraping websites data
@botbuster.post("/webscraper/") # endpoint #4 scraping data from websites
async def web_scraping(request: ds.web_scraper, background_tasks: BackgroundTasks):
    page_url = request.dict()["page_url"]
    if not validation.is_valid_url(page_url):
        raise HTTPException(status_code=400, detail="Bad Request")
    with open(CONFIG_FILE_PATH, "r") as config_data_file:
        website_configs = json.load(config_data_file)["website_configs"]
    website_name = urlparse(page_url).netloc
    scraping_data = website_configs.get(website_name, website_configs["default"])

    if scraping_data["func"] == "sms":
        items = await ws.scraper(scraping_data["elements"], scraping_data["settings"], page_url)
    elif scraping_data["func"] == "gws":
        url = scraping_data.get("url", None)
        splitter = scraping_data.get("splitter", None)
        items = ws.generic_scraper(scraping_data["elements"], page_url, url, splitter)
    else: 
        raise HTTPException(status_code=500, detail="Internal Server Error")

    return validation.sanitise(items)


# get website scraping configurations
@botbuster.get("/webscraper/settings/") # endpoint to retrieve the webscraping settings
async def get_webscraper_settings():
    with open(CONFIG_FILE_PATH, "r") as config_data_file:
        website_configs = json.load(config_data_file)["website_configs"]
    website_settings = {key: website_configs[key]["elements"] for key in website_configs.keys()}
    return website_settings
    
# save web scraping settings
@botbuster.post("/webscraper/settings/")
async def update_config(website_configs: dict):
    # try:
    # validation
    # for domain, selectors in website_configs.items():
    #     # if not validation.check_domain(domain):
    #     #     raise HTTPException(status_code=400, detail="Invalid Input!")
    #     for selector in selectors:
    #         if not validation.check_elements(selector):
    #             raise HTTPException(status_code=400, detail="Invalid Input!")
        # save the updated configuration to the config.json file
        with open(TEMP_FILE_PATH, "r") as f:
            config_data = json.load(f)
        # update the existing config file with the new data
        for key, value in website_configs.items():
            for website_url in config_data[key].keys():
                config_data[key][website_url]["elements"]= value[website_url]
            for website_url in value:
                if website_url not in config_data[key]:
                    config_data[key][website_url] = {"elements": value[website_url], "settings": ["",50],"func": "sms"}
            # save the updated configuration to the config.json file
            with open(TEMP_FILE_PATH, "w") as f:
                json.dump(config_data, f, indent=3)

#append new elements that user enters 
@botbuster.post("/webscraper/update_elements/")      
async def update_elements(request: Request):
    data = await request.json()
    website = data.get("website")
    new_element = data.get("newElement")
    # validation
    # if(validation.check_domain(website)==False):
    #     raise HTTPException(status_code = 400, detail = "Invalid Input!")
    # if(validation.check_elements(new_element)==False):
    #     raise HTTPException(status_code = 400, detail = "Invalid Input!")
    
    with open(TEMP_FILE_PATH, "r") as f:
        config_data = json.load(f)    
    if website in config_data["website_configs"]:
        config_data["website_configs"][website]["elements"].append(new_element)
    with open(TEMP_FILE_PATH, "w") as f:
            json.dump(config_data, f, indent=3)

# save the user's changes
@botbuster.get("/webscraper/temp_config")
def changeFile():
    with open(TEMP_FILE_PATH, "r") as f:
        contents = f.read()
    with open(CONFIG_FILE_PATH, "w") as f: #writing the temp file to config file
        f.write(contents)

# save the user's changes
@botbuster.get("/webscraper/del_changes")
def changeFile():
    with open(CONFIG_FILE_PATH, "r") as f:
        contents = f.read()
    with open(TEMP_FILE_PATH, "w") as f: #resetting the temp file if user did not press save
        f.write(contents)

# extracting text from user's file
@botbuster.post("/extract/") # endpoint #4 scraping data from files
def extract_text(file: UploadFile):
    try: 
        file_extension = file.filename.rsplit('.', 1)[1].lower() #extracts file extension from file name
        contents = file.file.read() # gets content of file out in bytes
        with open(f"temp.{file_extension}", 'wb') as f:
            f.write(contents) # writes a temporary file with the same bytes
        if file_extension == 'docx': # if it's docx file
            doc = docx.Document(f"temp.{file_extension}")
            section = doc.sections[0]
            header = section.header
            header.paragraphs[0].text = ""
            footer = section.footer
            footer.paragraphs[0].text = ""
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            text = '\n'.join(paragraphs)
        elif file_extension == 'txt':  # if it's txt file
            text = textract.process(f"temp.{file_extension}", method = "python").decode('utf-8')
        elif file_extension == 'pdf': # if it's pdf file    
            reader = PdfReader(f"temp.{file_extension}") # creating a pdf reader object
            text = ""
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                content = page.extract_text()
                lines = content.split('\n') 
                lines = lines[int(len(lines) * 0.05):int(len(lines))] #how much to take away (5%)
                cleaned_content = '\n'.join(lines)
                text += cleaned_content + '\n'
        elif file_extension == "jpg": # if it's image files
            image = np.array(Image.open(f"temp.{file_extension}")) # create image
            normalised_image = np.zeros((image.shape[0], image.shape[1]))
            # remove distractions in the photo using cv
            image = cv2.normalize(image, normalised_image, 0, 255, cv2.NORM_MINMAX)
            image = cv2.threshold(image, 100, 255, cv2.THRESH_BINARY)[1]
            image = cv2.GaussianBlur(image, (1, 1), 0)

            text = pytesseract.image_to_string(image) # use pytesseract to extract text
        else:
            return HTTPException(status_code = 400, detail = "Unsupported File Type")
    except:
        raise HTTPException(status_code = 500, detail = "Internal Server Error")
        
    finally:
        os.remove(f"temp.{file_extension}") # delete the temporary file whether there was an error or not
    return validation.sanitise(text)
    # return text